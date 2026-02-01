import os
from docx import Document
from docx.shared import Pt, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from typing import Dict, Any
from app.agent.graph import llm
from langchain_core.messages import SystemMessage, HumanMessage

class DocumentGenerator:
    def __init__(self):
        pass

    def generate_docx(self, doc_type: str, data: Dict[str, Any]) -> BytesIO:
        """
        Generate a .docx file based on types and user data.
        """
        doc = Document()
        
        # --- GLOBAL STYLES (Decree 30/2020/ND-CP) ---
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(14)
        
        # Set A4 margins
        section = doc.sections[0]
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)
        section.left_margin = Mm(30)
        section.right_margin = Mm(20)
        
        if doc_type == "don_khieu_nai_sa_thai":
            self._build_don_khieu_nai(doc, data)
        elif doc_type == "don_ly_hon_thuan_tinh":
            self._build_don_ly_hon(doc, data)
        elif doc_type == "don_khoi_kien":
            self._build_don_khoi_kien(doc, data)
        elif doc_type == "hop_dong_dat_coc":
            self._build_hop_dong_dat_coc(doc, data)
        elif doc_type == "giay_uy_quyen":
            self._build_giay_uy_quyen(doc, data)
        elif doc_type == "hop_dong_thue_nha":
            self._build_hop_dong_thue_nha(doc, data)
        else:
            doc.add_heading("VĂN BẢN PHÁP LÝ", 0)
            doc.add_paragraph(f"Loại văn bản: {doc_type}")
            for key, value in data.items():
                doc.add_paragraph(f"{key}: {value}")

        file_stream = BytesIO()
        doc.save(file_stream)
        file_stream.seek(0)
        return file_stream

    async def refine_content(self, field_name: str, content: str, doc_type: str) -> str:
        """
        Use AI to refine rough user notes into professional legal text.
        """
        system_prompt = f"""Bạn là một chuyên gia soạn thảo văn bản pháp luật tiêu chuẩn Việt Nam.
Nhiệm vụ: Chuyển đổi các ghi chú thô của người dùng thành ngôn ngữ pháp lý chuyên nghiệp, trang trọng và chính xác.
Loại văn bản: {doc_type}
Trường nội dung: {field_name}

QUY TẮC:
- Sử dụng thuật ngữ pháp lý chính xác (VD: "đơn phương chấm dứt", "quy định tại", "ký kết").
- Ngôn ngữ trang trọng, khách quan, không biểu cảm thái quá.
- Giữ nguyên ý nghĩa cốt lõi của người dùng nhưng diễn đạt lại mạch lạc.
- CHỈ TRẢ VỀ nội dung đã được tinh chỉnh, không kèm theo lời dẫn hay giải thích.
"""
        messages = [
            SystemMessage(content=system_prompt),
            HumanMessage(content=f"Ghi chú thô: {content}")
        ]
        
        try:
            response = await llm.ainvoke(messages)
            return response.content.strip()
        except Exception as e:
            print(f"Error refining content with AI: {e}")
            return content # Fallback to original content

    def generate_preview(self, doc_type: str, data: Dict[str, Any]) -> str:
        """
        Produce a high-quality text preview for the frontend.
        Matches the structure of the .docx generator.
        """
        header = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n---------------"
        footer_note = "\n\nTôi xin cam đoan nội dung trên là đúng sự thật.\nTrân trọng cảm ơn!"
        footer_sign = f"\n\n{' ' * 40}......, ngày ... tháng ... năm ...\n{' ' * 40}Người làm đơn/Người cam đoan\n{' ' * 45}(Ký và ghi rõ họ tên)"

        if doc_type == "don_khieu_nai_sa_thai":
            content = f"\n\nĐƠN KHIẾU NẠI\n(Về việc: {data.get('reason', 'Sa thải công nhân viên trái quy định')})\n\n"
            content += f"Kính gửi: {data.get('recipient', '[Tên cơ quan/Công ty/Người có thẩm quyền]')}\n\n"
            content += f"Tôi tên là: {data.get('full_name', '................................')}\n"
            content += f"Địa chỉ thường trú: {data.get('address', '................................')}\n\n"
            content += f"NỘI DUNG KHIẾU NẠI:\n{data.get('content', '................................')}\n\n"
            content += f"YÊU CẦU GIẢI QUYẾT:\n{data.get('request', '................................')}"
            return header + content + footer_note + footer_sign

        elif doc_type == "don_ly_hon_thuan_tinh":
            content = f"\n\nĐƠN YÊU CẦU CÔNG NHẬN THUẬN TÌNH LY HÔN,\nNUÔI CON, CHIA TÀI SẢN KHI LY HÔN\n\n"
            content += f"Kính gửi: Tòa án nhân dân {data.get('recipient', '....................')}\n\n"
            content += f"Chúng tôi gồm:\n"
            content += f"Vợ: {data.get('full_name', '....................')} - Sinh năm: ..........\n"
            content += f"Địa chỉ cư trú: {data.get('address', '....................')}\n"
            content += f"Chồng: ....................................... - Sinh năm: ..........\n"
            content += f"Địa chỉ cư trú: .......................................\n\n"
            content += f"NỘI DUNG YÊU CẦU:\n{data.get('content', '................................')}\n\n"
            content += f"VỀ CON CHUNG VÀ TÀI SẢN:\n{data.get('request', '................................')}"
            return header + content + footer_note + footer_sign

        elif doc_type == "don_khoi_kien":
            content = f"\n\nĐƠN KHỞI KIỆN\n\n"
            content += f"Kính gửi: Tòa án nhân dân {data.get('recipient', '....................')}\n\n"
            content += f"NGƯỜI KHỞI KIỆN:\n"
            content += f"Họ và tên: {data.get('full_name', '....................')}\n"
            content += f"Địa chỉ cư trú: {data.get('address', '....................')}\n\n"
            content += f"NGƯỜI BỊ KIỆN:\n"
            content += f"Họ và tên/Cơ quan: {data.get('recipient', '....................')}\n\n"
            content += f"NỘI DUNG KHỞI KIỆN:\n{data.get('content', '................................')}\n\n"
            content += f"YÊU CẦU TÒA ÁN GIẢI QUYẾT:\n{data.get('request', '................................')}"
            return header + content + footer_note + footer_sign

        elif doc_type == "hop_dong_dat_coc":
            content = f"\n\nHỢP ĐỒNG ĐẶT CỌC\n(Số: .../.../HĐĐC)\n\n"
            content += f"Hôm nay, ngày ... tháng ... năm ..., tại .......................................\n\n"
            content += f"BÊN ĐẶT CỌC (BÊN A):\n"
            content += f"Họ và tên: {data.get('full_name', '....................')} - CMND/CCCD: ....................\n"
            content += f"Địa chỉ: {data.get('address', '....................')}\n\n"
            content += f"BÊN NHẬN ĐẶT CỌC (BÊN B):\n"
            content += f"Họ và tên: {data.get('recipient', '....................')} - CMND/CCCD: ....................\n\n"
            content += f"NỘI DUNG THỎA THUẬN:\n"
            content += f"Bên A đồng ý đặt cọc cho bên B số tiền để bảo đảm thực hiện việc: {data.get('reason', '....................')}\n"
            content += f"Số tiền đặt cọc (bằng số): {data.get('request', '....................')}\n"
            content += f"Nội dung chi tiết: {data.get('content', '....................')}\n\n"
            content += f"Hợp đồng được lập thành 02 bản có giá trị pháp lý như nhau.\n\n"
            content += f"{' ' * 5}BÊN A (Ký tên) {' ' * 30} BÊN B (Ký tên)"
            return content

        elif doc_type == "giay_uy_quyen":
            content = f"\n\nGIẤY ỦY QUYỀN\n\n"
            content += f"NGƯỜI ỦY QUYỀN (BÊN A):\n"
            content += f"Họ và tên: {data.get('full_name', '....................')}\n"
            content += f"Địa chỉ thường trú: {data.get('address', '....................')}\n\n"
            content += f"NGƯỜI ĐƯỢC ỦY QUYỀN (BÊN B):\n"
            content += f"Họ và tên: {data.get('recipient', '....................')}\n\n"
            content += f"NỘI DUNG ỦY QUYỀN:\n"
            content += f"Phạm vi ủy quyền: {data.get('content', '....................')}\n"
            content += f"Thời hạn ủy quyền: {data.get('request', '....................')}"
            return header + content + footer_note + footer_sign

        elif doc_type == "hop_dong_thue_nha":
            content = f"\n\nHỢP ĐỒNG THUÊ NHÀ Ở\n\n"
            content += f"BÊN CHO THUÊ (BÊN A):\n"
            content += f"Họ và tên: {data.get('recipient', '....................')}\n\n"
            content += f"BÊN THUÊ (BÊN B):\n"
            content += f"Họ và tên: {data.get('full_name', '....................')}\n"
            content += f"Địa chỉ thường trú: {data.get('address', '....................')}\n\n"
            content += f"NỘI DUNG HỢP ĐỒNG:\n{data.get('content', '................................')}\n\n"
            content += f"GIÁ THUÊ VÀ PHƯƠNG THỨC THANH TOÁN:\n{data.get('request', '................................')}\n\n"
            content += f"{' ' * 5}BÊN A (Ký tên) {' ' * 30} BÊN B (Ký tên)"
            return content

        return "VĂN BẢN PHÁP LÝ\n\n" + "\n".join([f"{k}: {v}" for k, v in data.items()])

    def _build_don_khieu_nai(self, doc, data):
        self._add_legal_header(doc)
        
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\nĐƠN KHIẾU NẠI")
        run.bold = True
        run.font.size = Pt(16)
        
        subject = doc.add_paragraph()
        subject.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subject.add_run(f"(Về việc: {data.get('reason', 'Sa thải trái quy định')})").italic = True
        
        doc.add_paragraph(f"\nKính gửi: {data.get('recipient', '[Tên cơ quan/Công ty]')}")
        
        p = doc.add_paragraph()
        p.add_run("Tôi tên là: ").bold = True
        p.add_run(data.get('full_name', '................................'))
        
        p = doc.add_paragraph()
        p.add_run("Địa chỉ: ").bold = True
        p.add_run(data.get('address', '................................'))
        
        doc.add_paragraph("\nNỘI DUNG KHIẾU NẠI:").bold = True
        doc.add_paragraph(data.get('content', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nYÊU CẦU GIẢI QUYẾT:").bold = True
        doc.add_paragraph(data.get('request', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nTôi xin cam đoan nội dung trên là đúng sự thật.\nTrân trọng cảm ơn!")
        self._add_signature_area(doc, "Người làm đơn/Người khiếu nại")

    def _build_don_ly_hon(self, doc, data):
        self._add_legal_header(doc)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\nĐƠN YÊU CẦU CÔNG NHẬN THUẬN TÌNH LY HÔN\nNUÔI CON, CHIA TÀI SẢN KHI LY HÔN")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph(f"\nKính gửi: Tòa án nhân dân {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nChúng tôi gồm:").bold = True
        doc.add_paragraph(f"Vợ: {data.get('full_name', '....................')} - Địa chỉ: {data.get('address', '....................')}")
        doc.add_paragraph(f"Chồng: ....................................... - Địa chỉ: .......................................")
        
        doc.add_paragraph("\nNỘI DUNG YÊU CẦU:").bold = True
        doc.add_paragraph(data.get('content', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nVỀ CON CHUNG VÀ TÀI SẢN:").bold = True
        doc.add_paragraph(data.get('request', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self._add_signature_table_double(doc, "Vợ", "Chồng")

    def _build_don_khoi_kien(self, doc, data):
        self._add_legal_header(doc)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\nĐƠN KHỞI KIỆN")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph(f"\nKính gửi: Tòa án nhân dân {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nNGƯỜI KHỞI KIỆN:").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('full_name', '....................')} - Địa chỉ: {data.get('address', '....................')}")
        
        doc.add_paragraph("\nNGƯỜI BỊ KIỆN:").bold = True
        doc.add_paragraph(f"Họ tên/Cơ quan: {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nNỘI DUNG KHỞI KIỆN:").bold = True
        doc.add_paragraph(data.get('content', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nYÊU CẦU TÒA ÁN GIẢI QUYẾT:").bold = True
        doc.add_paragraph(data.get('request', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self._add_signature_area(doc, "Người khởi kiện")

    def _build_hop_dong_dat_coc(self, doc, data):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("HỢP ĐỒNG ĐẶT CỌC")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph("\nHôm nay, ngày ... tháng ... năm ..., tại .......................................")
        
        doc.add_paragraph("\nBÊN ĐẶT CỌC (BÊN A):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('full_name', '....................')} - Địa chỉ: {data.get('address', '....................')}")
        
        doc.add_paragraph("\nBÊN NHẬN ĐẶT CỌC (BÊN B):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nNỘI DUNG THỎA THUẬN:").bold = True
        doc.add_paragraph(f"Bên A đặt cọc cho bên B số tiền để bảo đảm thực hiện việc: {data.get('reason', '....................')}")
        doc.add_paragraph(f"Số tiền đặt cọc: {data.get('request', '....................')}")
        doc.add_paragraph(f"Nội dung chi tiết: {data.get('content', '....................')}")
        
        doc.add_paragraph("\nHợp đồng được lập thành 02 bản có giá trị pháp lý như nhau.")
        self._add_signature_table_double(doc, "BÊN A", "BÊN B")

    def _build_giay_uy_quyen(self, doc, data):
        self._add_legal_header(doc)
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("\nGIẤY ỦY QUYỀN")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph("\nNGƯỜI ỦY QUYỀN (BÊN A):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('full_name', '....................')} - Địa chỉ: {data.get('address', '....................')}")
        
        doc.add_paragraph("\nNGƯỜI ĐƯỢC ỦY QUYỀN (BÊN B):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nNỘI DUNG ỦY QUYỀN:").bold = True
        doc.add_paragraph(f"Phạm vi ủy quyền: {data.get('content', '....................')}")
        doc.add_paragraph(f"Thời hạn ủy quyền: {data.get('request', '....................')}")
        
        self._add_signature_table_double(doc, "BÊN A (Người ủy quyền)", "BÊN B (Người được ủy quyền)")

    def _build_hop_dong_thue_nha(self, doc, data):
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("HỢP ĐỒNG THUÊ NHÀ Ở")
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph("\nBÊN CHO THUÊ (BÊN A):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('recipient', '....................')}")
        
        doc.add_paragraph("\nBÊN THUÊ (BÊN B):").bold = True
        doc.add_paragraph(f"Họ tên: {data.get('full_name', '....................')} - Địa chỉ: {data.get('address', '....................')}")
        
        doc.add_paragraph("\nNỘI DUNG HỢP ĐỒNG:").bold = True
        doc.add_paragraph(data.get('content', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        doc.add_paragraph("\nGIÁ THUÊ VÀ PHƯƠNG THỨC THANH TOÁN:").bold = True
        doc.add_paragraph(data.get('request', '................................')).paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        self._add_signature_table_double(doc, "BÊN A", "BÊN B")

    def _add_legal_header(self, doc):
        """Unified National Header for all documents"""
        p1 = doc.add_paragraph()
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
        run1.bold = True
        run1.font.size = Pt(13)
        p1.paragraph_format.space_after = 0
        
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run2 = p2.add_run("Độc lập - Tự do - Hạnh phúc")
        run2.bold = True
        run2.font.size = Pt(14)
        p2.paragraph_format.space_after = 0
        
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run("---------------")

    def _add_signature_area(self, doc, role="Người làm đơn", date_line="......, ngày ... tháng ... năm ..."):
        """Professional Signature Block using an invisible table for perfect alignment"""
        doc.add_paragraph() # Spacer
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        # Cell 0: Left side (usually empty in single signature docs)
        # Cell 1: Right side signature
        cell = table.cell(0, 1)
        
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_date = p1.add_run(date_line)
        run_date.font.italic = True
        
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_role = p2.add_run(role)
        run_role.bold = True
        
        p3 = cell.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_sign = p3.add_run("(Ký và ghi rõ họ tên)")
        run_sign.font.italic = True

    def _add_signature_table_double(self, doc, role_a, role_b):
        """Standard double signature layout using invisible table"""
        doc.add_paragraph()
        table = doc.add_table(rows=1, cols=2)
        table.autofit = True
        
        c0 = table.cell(0, 0)
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run0 = p0.add_run(role_a)
        run0.bold = True
        c0.add_paragraph("(Ký và ghi rõ họ tên)").alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        c1 = table.cell(0, 1)
        p1 = c1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run1 = p1.add_run(role_b)
        run1.bold = True
        c1.add_paragraph("(Ký và ghi rõ họ tên)").alignment = WD_ALIGN_PARAGRAPH.CENTER

doc_generator = DocumentGenerator()
